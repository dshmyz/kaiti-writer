#!/usr/bin/env python3
"""GB/T 7714 顺序编码制引用编号校验与重排工具。

顺序编码制要求：正文引用标注 [n] 按首次出现顺序递增（[8] 不得早于 [2] 出现），
参考文献表按编号升序排列。插入新文献后重编号全靠手改是踩过的坑，本工具自动完成。

用法：
  校验（只报不修）：  python renumber_refs.py --docx 报告.docx --content content.json --check
  自动重排（正文+文献表同步）：  python renumber_refs.py --docx 报告.docx --content content.json --fix
  插入新文献：content.json 的 refs 数组末尾加 "[99] 新文献条目"，正文引用处写 [99]，
  然后跑 --fix，99 及其后所有编号自动归位（99→实际位次，原 19-21 顺延为 20-22）。

原理：扫描正文引用标注的首次出现顺序 → 建立 旧编号→新编号 映射 → 同步重写
正文标注与文献表序号。--check 与 build_and_validate.py 第 4.5 步规则一致。
"""
import argparse
import json
import re
import sys


CITE_RE = re.compile(r"\[(\d+(?:\s*[,，]\s*\d+)*(?:\s*-\s*\d+)?)\]")


def extract_first_order(docx_path):
    """从 docx 正文（参考文献表之前）提取引用编号的首次出现顺序。"""
    try:
        from docx import Document
    except ImportError:
        sys.exit("缺少 python-docx：pip install python-docx")
    doc = Document(docx_path)
    refs_start = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        # 取最后一个"参考文献"标题：目录里的"参考文献\t页码"占位在前，
        # 正文插入的才是真内容（与 build_and_validate.py 第 4 节同规则）
        ref_title = t.split(chr(9))[0] if chr(9) in t else t
        if "参考文献" in ref_title and len(ref_title) < 10:
            refs_start = i
    order, seen = [], set()
    for i, p in enumerate(doc.paragraphs):
        if refs_start is not None and i == refs_start:
            break
        for m in CITE_RE.finditer(p.text):
            for num_str in re.split(r"[,，\-]", m.group(1)):
                n = int(num_str.strip())
                if n and n not in seen:
                    seen.add(n)
                    order.append(n)
    return order, doc


def build_mapping(order):
    """首次出现顺序 → 新编号映射（首次出现的排 1, 2, 3…）。"""
    return {old: new + 1 for new, old in enumerate(order)}


def remap_cite(match, mapping):
    """把一处标注 [3, 5-7] 的每个编号按映射替换，保持多编号/区间结构。"""
    inner = match.group(1)
    def sub_num(m):
        old = int(m.group(0))
        return str(mapping.get(old, old))
    return "[" + re.sub(r"\d+", sub_num, inner) + "]"


def sort_citation(text):
    """重排后标注内的编号可能乱序（如 [5, 2]），排序回 [2, 5]。"""
    def _sort(m):
        nums = [int(x) for x in re.split(r"[,，\s]+", m.group(1).strip()) if x]
        if len(nums) > 1 and nums != sorted(nums):
            return "[" + ", ".join(str(n) for n in sorted(nums)) + "]"
        return m.group(0)
    return re.sub(r"\[(\d+(?:\s*[,，]\s*\d+)+)\]", _sort, text)


def check_docx(order):
    """返回乱序编号列表（空 = 合规）。"""
    return [b for a, b in zip(order, order[1:]) if b < a]


def _replace_once_in_sections(data, old_s, new_s):
    """在 content_by_section 的所有文本块里做替换，返回命中次数（不写回重复命中）。"""
    hit = 0
    for blocks in data.get("content_by_section", {}).values():
        if not isinstance(blocks, list):
            continue
        for i, blk in enumerate(blocks):
            if isinstance(blk, str) and old_s in blk:
                blocks[i] = blk.replace(old_s, new_s)
                hit += 1
    return hit


def main():
    ap = argparse.ArgumentParser(description="GB/T 7714 顺序编码制引用编号校验与重排")
    ap.add_argument("--docx", required=True, help="开题报告 .docx 路径")
    ap.add_argument("--content", help="content.json 路径（--fix 时同步重排 refs 序号）")
    ap.add_argument("--check", action="store_true", help="只校验不修改")
    ap.add_argument("--fix", action="store_true", help="自动重排正文标注 + 文献表序号")
    ap.add_argument("--insert", default=None,
                    help='"旧文本=新文本;..." 先对 content.json 正文做唯一性替换再重排'
                         "（插入新文献：refs 末尾加 [99] 占位条目，正文旧句子替换为含 [99] 的新句子，跑 --fix 归位）")
    args = ap.parse_args()

    if not args.check and not args.fix and not args.insert:
        ap.error("至少指定 --check 或 --fix 之一（--insert 可单独用，等价于替换正文后引导 rebuild+fix）")

    # --- --insert：先替换 content.json 正文（唯一性校验，WorkBuddy 版优点） ---
    if args.insert:
        if not args.content:
            ap.error("--insert 需要同时指定 --content")
        with open(args.content, encoding="utf-8") as f:
            data0 = json.load(f)
        n_ins = 0
        for pair in args.insert.split(";"):
            if "=" not in pair:
                print(f"⚠️  --insert 片段缺少 '='：{pair[:40]}")
                sys.exit(1)
            old_s, new_s = pair.split("=", 1)
            hit = _replace_once_in_sections(data0, old_s, new_s)
            if hit != 1:
                print(f"⚠️  插入失败：\"{old_s[:30]}...\" 命中 {hit} 次（应为 1，旧文本须唯一）")
                sys.exit(1)
            n_ins += hit
        with open(args.content, "w", encoding="utf-8") as f:
            json.dump(data0, f, ensure_ascii=False, indent=2)
        print(f"✅ 已替换 {n_ins} 处正文（content.json）")
        print("⚠️  请先重新 build docx（让正文含新引用），再跑一次 --fix 完成重排：")
        print("    python build_and_validate.py --template ... --content ... --output ...")
        print("    python renumber_refs.py --docx ... --content ... --fix")
        return

    order, doc = extract_first_order(args.docx)
    if not order:
        print("正文无引用标注，无需处理")
        return
    # 友好检查：正文引用的编号是否都在文献表里（WorkBuddy 版优点——
    # 引用了不存在的编号时给明确指引，而不是重排出一堆错号）
    ref_nums = set()
    for p in doc.paragraphs:
        m0 = re.match(r"^\[(\d+)\]", p.text.strip())
        if m0:
            ref_nums.add(int(m0.group(1)))
    ghost = [n for n in order if n not in ref_nums]
    if ghost:
        print(f"⚠️  正文引用了文献表中不存在的编号 {ghost[:8]}：")
        print("    请先在 content.json 的 refs 中补齐这些条目（编号随意，重排后会自动修正），")
        print("    或检查是否是手滑写错的引用号，再重新运行本工具。")
        sys.exit(1)
    bad = check_docx(order)
    if not bad and sorted(order) == list(range(1, len(order) + 1)):
        print(f"✅ 引用编号按首次出现顺序合规（{min(order)}-{max(order)}，共 {len(order)} 个编号）")
        return
    if args.check:
        if not bad and sorted(order) != list(range(1, len(order) + 1)):
            print(f"⚠️  首次出现顺序无乱序，但编号不连续：{sorted(order)}（如 [99] 占位未归位）")
        else:
            print(f"⚠️  引用编号未按首次出现顺序：乱序出现在 {bad[:8]}")
        print("    → 运行本工具 --fix 自动重排")
        sys.exit(1)
        print(f"⚠️  引用编号未按首次出现顺序：乱序出现在 {bad[:8]}")
        print("    → 运行本工具 --fix 自动重排")
        sys.exit(1)

    # --- fix ---
    mapping = build_mapping(order)
    print(f"重排映射（旧→新）：{ {k: v for k, v in sorted(mapping.items()) if k != v} or '无需变动' }")

    # 1) 重写 docx 正文标注
    n_cites = 0
    refs_start = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        ref_title = t.split(chr(9))[0] if chr(9) in t else t
        if "参考文献" in ref_title and len(ref_title) < 10:
            refs_start = i  # 不 break：取最后一个（目录占位在前）
    for i, p in enumerate(doc.paragraphs):
        if refs_start is not None and i == refs_start:
            break
        for run in p.runs:
            if run.text and CITE_RE.search(run.text):
                run.text = sort_citation(CITE_RE.sub(lambda m: remap_cite(m, mapping), run.text))
                n_cites += 1

    # 2) 重写文献表序号（[旧] → [新]，按映射）
    if refs_start is not None:
        for p in doc.paragraphs[refs_start + 1:]:
            t = p.text.strip()
            if t.startswith(("附录", "附件")):
                break
            m = re.match(r"^\[(\d+)\]", t)
            if not m:
                continue
            old = int(m.group(1))
            new = mapping.get(old)
            if new and new != old:
                for run in p.runs:
                    if run.text and f"[{old}]" in run.text:
                        run.text = run.text.replace(f"[{old}]", f"[{new}]", 1)
                        break
                else:
                    # 序号可能在首个 run 与文本分离时——整体重写首 run
                    if p.runs:
                        p.runs[0].text = re.sub(r"^\[\d+\]", f"[{new}]", p.runs[0].text)
    doc.save(args.docx)
    print(f"✅ 已重排：正文 {n_cites} 处标注 + 文献表序号，保存至 {args.docx}")

    # 3) 同步重排 content.json：refs 数组（按新编号排序）+ 正文文本里的引用标注
    if args.content:
        with open(args.content, encoding="utf-8") as f:
            data = json.load(f)

        # 3a) 所有文本字段里的 [n] 标注按映射重写（正文引用在 section 文本里）
        # 注意：refs 数组单独处理（文献表序号不能被 3a 当正文标注改掉，否则撞号）
        refs_backup = data.get("refs", [])
        data.pop("refs", None)

        def remap_text(v):
            if isinstance(v, str):
                if CITE_RE.search(v):
                    return sort_citation(CITE_RE.sub(lambda m: remap_cite(m, mapping), v))
                return v
            if isinstance(v, list):
                return [remap_text(x) for x in v]
            if isinstance(v, dict):
                return {k: remap_text(x) for k, x in v.items()}
            return v
        data = remap_text(data)
        data["refs"] = refs_backup

        refs = data.get("refs", [])
        # 补全映射：refs 里正文未引用的编号（如占位条目、或多条文献共用正文一处引用），
        # 按其在 refs 数组中的出现顺序接在已用编号之后，避免撞号/断号
        full_mapping = dict(mapping)
        used = set(mapping.values())
        next_n = max(used, default=0) + 1
        for s in refs:
            m = re.match(r"^\[(\d+)\]", s)
            if not m:
                continue
            old_n = int(m.group(1))
            if old_n not in full_mapping:
                full_mapping[old_n] = next_n
                next_n += 1
        data["refs"] = sorted(refs, key=lambda s: full_mapping.get(
            int(m.group(1)), 10**9) if (m := re.match(r"^\[(\d+)\]", s)) else 10**9)
        fixed = []
        for s in data["refs"]:
            m = re.match(r"^\[(\d+)\]", s)
            if m:
                old_n = int(m.group(1))
                s = re.sub(r"^\[\d+\]", f"[{full_mapping[old_n]}]", s)
            fixed.append(s)
        data["refs"] = fixed
        with open(args.content, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"✅ content.json refs 已同步重排（{len(fixed)} 条）")


if __name__ == "__main__":
    main()
