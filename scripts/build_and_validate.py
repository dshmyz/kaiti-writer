#!/usr/bin/env python3
"""
一站式构建 + 校验脚本
用法：python build_and_validate.py --template 模板.docx --content content.json --output 输出.docx

流程：
  1. build_from_template.py — 从 content.json 生成 .docx
  2. fix_refs.py — 参考文献收尾核查（自动修正编号/年份/作者）
  3. 规范检查 — outlineLvl/盘古空格/附录颜色/参考文献表/字体字号/页面设置
"""
import argparse
import os
import re
import sys
import subprocess


# 字号 → pt 映射
FONT_SIZE_MAP = {
    "初号": 42, "小初": 36,
    "一号": 26, "小一": 24,
    "二号": 22, "小二": 18,
    "三号": 16, "小三": 15,
    "四号": 14, "小四": 12,
    "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5,
    "七号": 5.5, "八号": 5,
}

# 格式规范.md 的字体字号要求（章/节/条/正文）
HEADING_SPECS = {
    # outlineLvl: (expected_font, expected_size_pt, expected_alignment)
    0: ("黑体", 14, "center"),    # 章：黑体居中（模板14pt）
    1: ("黑体", 14, "left"),      # 节：四号黑体居左
    2: ("黑体", 12, "left"),      # 条：小四号黑体居左
}
BODY_SPEC = ("宋体", 12)  # 正文：小四号宋体
PAGE_NUM_SPEC = ("宋体", 10.5)  # 页码：五号宋体


def run_step(name, cmd):
    """运行一个步骤，成功返回 True，失败返回 False。"""
    print(f"\n{'='*50}")
    print(f"  {name}")
    print(f"{'='*50}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.stdout:
        print(result.stdout)
    if result.returncode != 0:
        print(f"❌ {name} 失败：", file=sys.stderr)
        if result.stderr:
            print(result.stderr, file=sys.stderr)
        return False
    return True


def get_east_asia_font(run):
    """从 run 的 rPr 中提取东亚字体名（rFonts 的 eastAsia 属性）。"""
    from docx.oxml.ns import qn
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return None
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        return None
    return rFonts.get(qn("w:eastAsia"))


def validate_docx(docx_path):
    """规范检查：outlineLvl/盘古空格/附录颜色/参考文献表/字体字号/页面设置。"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("⚠️  未安装 python-docx，跳过规范检查（pip install python-docx）")
        return True

    print(f"\n{'='*50}")
    print("  规范检查")
    print(f"{'='*50}")

    doc = Document(docx_path)
    issues = []
    fixed = 0

    # --- 1. outlineLvl ---
    headings = [p for p in doc.paragraphs
                if p.style and p.style.name and p.style.name.startswith("Heading")]
    for p in headings:
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            continue
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is not None:
            val = int(outline.get(qn("w:val")))
            if val > 2:
                issues.append(f"⚠️  大纲层级异常：「{p.text[:20]}」outlineLvl={val}")
    if not any("outlineLvl" in i for i in issues):
        print("✅ outlineLvl 全部正确")
    else:
        for i in issues:
            if "outlineLvl" in i:
                print(i)

    # --- 2. 盘古空格 ---
    # 只处理半角空格（" "），不用 \s——\s 会吞目录页码的制表符。
    # 跳过：标题段（Heading/样式/章标题/附录开头）、图表题注（图/表+数字开头）、
    #       参考文献条目（[数字]开头）——题注"图1␣␣XXX"的 2 格、
    #       GB/T 7714 标点后空格都是格式要求，不能动。
    def strip_halfwidth_cjk_spaces(text):
        t = re.sub(r"([一-鿿]) +([A-Za-z0-9(（])", r"\1\2", text)
        t = re.sub(r"([A-Za-z0-9)%）]) +([一-鿿])", r"\1\2", t)
        return t

    def para_skip_pangu(p):
        t = p.text.strip()
        if not t:
            return True
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            return True
        if pPr := p._element.find(qn("w:pPr")):
            if pPr.find(qn("w:outlineLvl")) is not None:
                return True
        if re.match(r"^(第[一二三四五六七八九十百]+[章节]|[一二三四五六七八九十]+、|附录|附件)", t):
            return True
        if re.match(r"^[图表]\s*\d+", t):
            return True
        if re.match(r"^\[\d+\]", t):
            return True
        return False

    cjk_space_count = 0
    for p in doc.paragraphs:
        if para_skip_pangu(p):
            continue
        for run in p.runs:
            if run.text:
                original = run.text
                t = strip_halfwidth_cjk_spaces(run.text)
                if t != original:
                    run.text = t
                    cjk_space_count += 1
    # 表格也要检查
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if para_skip_pangu(p):
                        continue
                    for run in p.runs:
                        if run.text:
                            original = run.text
                            t = strip_halfwidth_cjk_spaces(run.text)
                            if t != original:
                                run.text = t
                                cjk_space_count += 1
    if cjk_space_count:
        print(f"⚠️  发现 {cjk_space_count} 处盘古空格，已自动修正")
        fixed += cjk_space_count
    else:
        print("✅ 盘古空格无异常")

    # --- 3. 附录标题颜色 ---
    appendix_fixed = False
    for p in doc.paragraphs:
        if p.text.strip().startswith("附录"):
            for run in p.runs:
                rPr = run._element.find(qn("w:rPr"))
                if rPr is not None:
                    color = rPr.find(qn("w:color"))
                    if color is not None:
                        rgb = color.get(qn("w:val"), "")
                        if rgb.upper() == "FF0000":
                            rPr.remove(color)
                            appendix_fixed = True
                            fixed += 1
    if appendix_fixed:
        print("⚠️  附录标题为红色，已自动修正")
    else:
        print("✅ 附录标题颜色正常")

    # --- 4. 参考文献表 ---
    # 找最后一个"参考文献"标题（模板原位有一个带tab的占位，正文插入的才是真内容）
    refs_start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        ref_title = text.split(chr(9))[0] if chr(9) in text else text
        if "参考文献" in ref_title and len(ref_title) < 10:
            refs_start_idx = i
    in_refs = False
    ref_issues = []
    ref_count = 0
    expected_seq = 1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if i == refs_start_idx:
            in_refs = True
            continue
        if in_refs:
            # 遇到附录或下一个主要章节则停止
            if text and (text.startswith("附录") or text.startswith("附件")):
                break
            if not text:
                continue
            ref_count += 1
            m = re.match(r"^\[(\d+)\]", text)
            if not m:
                ref_issues.append(f"⚠️  第 {ref_count} 条缺少 [序号] 标记")
            else:
                seq = int(m.group(1))
                if seq != expected_seq:
                    ref_issues.append(f"⚠️  序号不连续：期望 [{expected_seq}] 实际 [{seq}]")
                expected_seq = seq + 1
            # GB/T 7714：每条末不加结束符（与 fix_refs.py、参考文献著录.md 一致）
            if text.endswith((".", "．", "。", ";", "；")):
                ref_issues.append(f"⚠️  第 {ref_count} 条末尾不应有结束符（GB/T 7714）")
    if ref_count == 0:
        print("⚠️  未找到参考文献表")
    elif ref_issues:
        for i in ref_issues[:5]:
            print(i)
        if len(ref_issues) > 5:
            print(f"  ...共 {len(ref_issues)} 条问题")
    else:
        print(f"✅ 参考文献表格式正常（共 {ref_count} 条）")

    # --- 4.5 引用编号按首次出现顺序（GB/T 7714 顺序编码制） ---
    # 正文引用标注 [n] 应按首次出现顺序递增：[8] 不得早于 [2] 出现。
    # 违规时提示运行 scripts/renumber_refs.py --fix 自动重排。
    cite_order = []
    seen = set()
    for i, p in enumerate(doc.paragraphs):
        if i == refs_start_idx:
            break  # 只扫正文，不扫参考文献表
        for m in re.finditer(r"\[(\d+(?:[,，]\s*\d+)*(?:-\d+)?)\]", p.text):
            for num_str in re.split(r"[,，\-]", m.group(1)):
                n = int(num_str.strip())
                if n and n not in seen:
                    seen.add(n)
                    cite_order.append(n)
    out_of_order = [b for a, b in zip(cite_order, cite_order[1:]) if b < a]
    if not cite_order:
        print("✅ 正文无引用标注（跳过顺序检查）")
    elif out_of_order:
        print(f"⚠️  引用编号未按首次出现顺序：出现顺序 {[min(cite_order), max(cite_order)]} 区间内有乱序 {out_of_order[:5]}")
        print("    → 运行 python \"$SKILL/scripts/renumber_refs.py\" --fix 自动重排")
    else:
        print(f"✅ 引用编号按首次出现顺序合规（1-{max(cite_order)}）")

    # --- 5. 字体与字号（按格式规范.md） ---
    font_issues = []
    checked = 0
    in_body = False  # 跳过封面，从第一个 Heading 或目录开始检查
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        # 检测是否进入正文区域（第一个 Heading 或 "目录" 出现后）
        if not in_body:
            if p.style and p.style.name and p.style.name.startswith("Heading"):
                in_body = True
            elif "目录" in p.text:
                in_body = True
            else:
                continue  # 封面内容，跳过
        # 判断是否为 Heading
        is_heading = p.style and p.style.name and p.style.name.startswith("Heading")
        if is_heading:
            pPr = p._element.find(qn("w:pPr"))
            outline_val = None
            if pPr is not None:
                outline = pPr.find(qn("w:outlineLvl"))
                if outline is not None:
                    outline_val = int(outline.get(qn("w:val")))
            if outline_val in HEADING_SPECS:
                expected_font, expected_size, expected_align = HEADING_SPECS[outline_val]
                for run in p.runs:
                    if run.text.strip():
                        ea = get_east_asia_font(run)
                        if ea and ea != expected_font and "Heading" not in ea:
                            font_issues.append(
                                f"⚠️  「{p.text[:15]}」字体应为{expected_font}，实际为{ea}")
                        if run.font.size:
                            actual_pt = run.font.size.pt
                            if abs(actual_pt - expected_size) > 0.5:
                                font_issues.append(
                                    f"⚠️  「{p.text[:15]}」字号应为{expected_size}pt，实际为{actual_pt}pt")
                        checked += 1
                        break  # 只查第一个 run
        else:
            # 正文段落：跳过一切"有 outlineLvl 的标题段"（节标题/参考文献标题，
            # 14pt 四号）+ 附录标题 + 图表题注（图/表+数字开头，五号）——
            # 都不是正文的 12pt。注意必须用当前段落的文本，
            # 不能引用上一个循环残留的变量（陈旧变量 bug 曾让检查只覆盖 5 段）
            pPr0 = p._element.find(qn("w:pPr"))
            if pPr0 is not None and pPr0.find(qn("w:outlineLvl")) is not None:
                continue
            ptext = p.text.strip()
            if ptext.startswith(("（", "附录", "附件")) or re.match(r"^[图表]\s*\d+", ptext):
                continue
            for run in p.runs:
                if run.text.strip():
                    if run.font.size:
                        actual_pt = run.font.size.pt
                        if abs(actual_pt - BODY_SPEC[1]) > 1:
                            font_issues.append(
                                f"⚠️  正文字号应为{BODY_SPEC[1]}pt，实际为{actual_pt}pt  「{ptext[:15]}」")
                    checked += 1
                    break
    if font_issues:
        for i in font_issues[:5]:
            print(i)
        if len(font_issues) > 5:
            print(f"  ...共 {len(font_issues)} 条问题")
    else:
        print(f"✅ 字体字号符合要求（检查了 {checked} 个段落）")

    # --- 6. 页面设置（格式规范.md：上下左右各 25mm） ---
    for section in doc.sections:
        top = section.top_margin
        bottom = section.bottom_margin
        left = section.left_margin
        right = section.right_margin
        mm = 914400 / 25.4  # 1mm = 914400/25.4 EMUs
        tolerance = 1 * mm  # ±1mm 误差
        expected = 25 * mm
        ok = True
        for name, val in [("上", top), ("下", bottom), ("左", left), ("右", right)]:
            if abs(val - expected) > tolerance:
                ok = False
        if ok:
            print("✅ 页面设置符合要求（上下左右各 25mm）")
        else:
            print(f"⚠️  页边距：上{top/mm:.0f}mm 下{bottom/mm:.0f}mm "
                  f"左{left/mm:.0f}mm 右{right/mm:.0f}mm（应为各 25mm）")
        break  # 只检查第一节

    # --- 7. 页码字号（五号宋体 10.5pt） ---
    page_num_issues = []
    for section in doc.sections:
        footer = section.footer
        if footer:
            for p in footer.paragraphs:
                for run in p.runs:
                    if run.text.strip():
                        if run.font.size and abs(run.font.size.pt - 10.5) > 0.5:
                            page_num_issues.append(
                                f"⚠️  页码字号应为10.5pt，实际为{run.font.size.pt}pt")
                        break
        break  # 只检查第一节
    if page_num_issues:
        for i in page_num_issues:
            print(i)
    else:
        print("✅ 页码字号符合要求")

    # --- 保存修正 ---
    if fixed:
        doc.save(docx_path)
        print(f"\n共自动修正 {fixed} 项，已重新保存 {docx_path}")
    else:
        print("\n✅ 无需修正")

    if issues or font_issues or ref_issues:
        return False
    return True


def main():
    parser = argparse.ArgumentParser(
        description="一站式构建 + 校验：build → fix_refs → 规范检查")
    parser.add_argument("--template", required=True, help="模板 .docx 路径")
    parser.add_argument("--content", required=True, help="content.json 路径")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    parser.add_argument("--diff", default=None, help="改稿时传入旧版 .docx 路径")
    parser.add_argument("--mark-missing", default=None,
                        help="给参考文献缺项条目加红色提示（序号:提示文案，分号分隔）")
    parser.add_argument("--unmark-red", action="store_true",
                        help="定稿：清除参考文献区所有红色待补提示")
    parser.add_argument("--clear-noise-red", action="store_true",
                        help="清全文档非确认红（正文/目录/附录），只保留含〈待确认〉标记的红")
    parser.add_argument("--skip-fix-refs", action="store_true", help="跳过 fix_refs 步骤")
    parser.add_argument("--skip-validate", action="store_true", help="跳过规范检查步骤")
    args = parser.parse_args()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    build_script = os.path.join(script_dir, "build_from_template.py")
    fix_script = os.path.join(script_dir, "fix_refs.py")

    # === Step 1: build ===
    build_cmd = [sys.executable, build_script,
                 "--template", args.template,
                 "--content", args.content,
                 "--output", args.output]
    if args.diff:
        build_cmd += ["--diff", args.diff]

    if not run_step("Step 1: 生成 .docx", build_cmd):
        sys.exit(1)

    # === Step 2: fix_refs ===
    if not args.skip_fix_refs:
        fix_cmd = [sys.executable, fix_script,
                   "--content", args.content,
                   "--docx", args.output]
        if not run_step("Step 2: 参考文献收尾核查", fix_cmd):
            print("⚠️  fix_refs 出错，继续执行...")

    # === Step 2.5: 标红/清红（fix_refs 的交互模式） ===
    if args.mark_missing:
        fix_cmd = [sys.executable, fix_script,
                   "--docx", args.output,
                   "--mark-missing", args.mark_missing]
        run_step("Step 2.5: 参考文献缺项标红", fix_cmd)

    if args.unmark_red:
        fix_cmd = [sys.executable, fix_script,
                   "--docx", args.output,
                   "--unmark-red"]
        run_step("Step 2.5: 清除参考文献区红色提示", fix_cmd)

    if args.clear_noise_red:
        fix_cmd = [sys.executable, fix_script,
                   "--docx", args.output,
                   "--clear-noise-red"]
        run_step("Step 2.5: 清除全文档噪声红", fix_cmd)

    # === Step 3: 规范检查 ===
    if not args.skip_validate:
        validate_docx(args.output)

    print(f"\n{'='*50}")
    print(f"  ✅ 全部完成：{args.output}")
    print(f"{'='*50}")


if __name__ == "__main__":
    main()
